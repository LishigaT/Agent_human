import streamlit as st
from docx import Document
from io import BytesIO
import google.generativeai as genai
import json, os, glob, re
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np

st.set_page_config(page_title="ADGM Corporate Agent – Compliance Checker", layout="wide")
st.title("ADGM Corporate Agent")
st.write("Upload `.docx` files to check compliance.")

#API Key
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=GEMINI_API_KEY)
except:
    st.warning("Gemini API key missing. Add `GEMINI_API_KEY` in Streamlit secrets to enable AI checks.")
    GEMINI_API_KEY = None

REQUIRED_DOCS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors"
    ]
}

def extract_text_from_docx(file_obj):
    file_obj.seek(0)
    doc = Document(file_obj)
    return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

def detect_process(filenames):
    names = " ".join(filenames).lower()
    if any(word in names for word in ["articles", "memorandum", "incorporation", "aoa", "moa"]):
        return "Company Incorporation"
    return "Unknown"

def load_reference_texts(ref_dir="references"):
    refs = []
    if not os.path.exists(ref_dir):
        return refs
    for path in glob.glob(os.path.join(ref_dir, "*.txt")):
        try:
            with open(path, "r", encoding="utf-8") as f:
                refs.append((os.path.basename(path), f.read()))
        except:
            pass
    return refs

def chunk_text(text, size=400, overlap=50):
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunks.append(" ".join(words[i:i+size]))
        i += size - overlap
    return chunks

def build_rag_context(doc_text, refs, top_k=3):
    chunks, sources = [], []
    for fname, text in refs:
        for i, chunk in enumerate(chunk_text(text, 200, 40)):
            sources.append(f"{fname}::chunk{i+1}")
            chunks.append(chunk)
    if not chunks: return ""
    try:
        vec = TfidfVectorizer(stop_words="english", max_features=20000)
        X = vec.fit_transform(chunks + [doc_text])
        sims = (X[:-1] @ X[-1].T).toarray().ravel()
        top_idx = np.argsort(sims)[-top_k:][::-1]
        return "\n\n".join(
            f"--- Source: {sources[i]} ---\n{chunks[i]}"
            for i in top_idx if sims[i] > 0
        )
    except:
        return ""

def normalize_snippet(sn):
    return re.sub(r"\s+", " ", sn.strip()).lower()

def find_paragraph_indices_containing(doc, snippet):
    sn = normalize_snippet(snippet)
    return [i for i, p in enumerate(doc.paragraphs) if sn in normalize_snippet(p.text)]

def annotate_docx_bytes(original_bytesio, issues):
    original_bytesio.seek(0)
    doc = Document(original_bytesio)
    for issue in issues:
        snippet = issue.get("snippet") or issue.get("section") or issue.get("text") or issue.get("issue", "")
        if not snippet: continue
        para_idx = find_paragraph_indices_containing(doc, snippet)
        if not para_idx:  # fuzzy match
            words = [w for w in re.findall(r"\w{4,}", snippet)][:6]
            for i, p in enumerate(doc.paragraphs):
                if any(w.lower() in normalize_snippet(p.text) for w in words):
                    para_idx.append(i); break
        for i in para_idx:
            run = doc.paragraphs[i].add_run(
                f" [ADGM Review | Severity: {issue.get('severity', 'Review')} | Suggestion: {issue.get('suggestion', '')}]"
            )
            run.italic = True
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

st.subheader("Upload Documents")
uploaded_files = st.file_uploader("Upload `.docx` files", type=["docx"], accept_multiple_files=True)

refs = load_reference_texts("references")
st.info(f"Loaded {len(refs)} reference file(s) from `references/`.")

if uploaded_files:
    filenames = [f.name for f in uploaded_files]
    st.write("Uploaded:", filenames)

    process = detect_process(filenames)
    st.write("Detected process:", process)

    missing_docs = []
    if process in REQUIRED_DOCS:
        missing_docs = [d for d in REQUIRED_DOCS[process] if d.lower() not in " ".join(filenames).lower()]
        if missing_docs:
            st.error(f"Missing: {', '.join(missing_docs)}")
        else:
            st.success("All required documents present.")
    else:
        st.info("Checklist skipped for unknown process.")

    combined_texts, combined_all = {}, ""
    for f in uploaded_files:
        try:
            bytes_data = f.getvalue()
            text = extract_text_from_docx(BytesIO(bytes_data))
            combined_texts[f.name] = {"text": text, "bytes": bytes_data}
            combined_all += f"\n\n--- {f.name} ---\n\n{text}"
        except Exception as e:
            st.error(f"Error reading {f.name}: {e}")

    st.subheader("Extracted Text (preview)")
    st.text_area("Combined text", combined_all[:10000], height=300)

    rag_context = build_rag_context(combined_all, refs)
    if rag_context:
        st.subheader("Relevant ADGM References")
        st.text_area("RAG context", rag_context[:15000], height=200)
    else:
        st.info("No reference matches found.")

    if GEMINI_API_KEY:
        if st.button("Run Compliance Check"):
            with st.spinner("Analyzing with Gemini..."):
                prompt = "\n".join([
                    "You are an ADGM Corporate Agent compliance assistant.",
                    "Analyze provided documents against ADGM references.",
                    "Return ONLY valid JSON with 'issues'.",
                    "=== References ===", rag_context,
                    "=== Documents ==="
                ] + [f"--- {n} ---\n{info['text'][:8000]}" for n, info in combined_texts.items()])

                try:
                    model = genai.GenerativeModel("gemini-1.5-flash")
                    resp = model.generate_content(prompt)
                    raw = resp.text.strip()
                    try:
                        data = json.loads(raw)
                    except:
                        match = re.search(r"(\{[\s\S]*\})", raw)
                        data = json.loads(match.group(1)) if match else {"issues": []}
                    issues = data.get("issues", [])
                    st.subheader("Findings")
                    st.write(issues if issues else "No issues found.")
                    if not issues: st.text_area("Raw AI Output", raw, height=300)
                except Exception as e:
                    st.error(f"AI error: {e}")
                    issues = []

                final_report = {
                    "process": process,
                    "documents_uploaded": len(filenames),
                    "missing_documents": missing_docs,
                    "issues_found": issues
                }

                st.subheader("Download Reports")
                st.download_button("JSON Report", json.dumps(final_report, indent=2),
                                   file_name="compliance_report.json", mime="application/json")

                issues_by_doc = {}
                for it in issues:
                    issues_by_doc.setdefault(it.get("document", ""), []).append(it)

                for name, info in combined_texts.items():
                    orig = BytesIO(info["bytes"])
                    ann = annotate_docx_bytes(orig, issues_by_doc.get(name, []))
                    st.download_button(f"Original: {name}", info["bytes"], file_name=f"original_{name}")
                    st.download_button(f"Annotated: reviewed_{name}", ann.getvalue(), file_name=f"reviewed_{name}")

                st.success("Done.")
    else:
        st.warning("Gemini key missing: AI checks disabled.")
